"""--------------1-------------------"""
from docx import Document

# Load the field report
doc = Document('ringing_protocol_2024.docx')

full_text = []
for para in doc.paragraphs:
    # Extract only text, ignore style for now
    if "Species:" in para.text:
        full_text.append(para.text)

print(f"Extracted {len(full_text)} lines.")

"""--------------2-------------------"""
from docx import Document

doc = Document() # Creates new blank doc
doc.add_heading('Annual Ringing Report', 0)

# Add a table with data
data = [('A12', 'Blaumeise'), ('B99', 'Rotkehlchen')]
table = doc.add_table(rows=1, cols=2)

# Set header
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Ring ID'
hdr_cells[1].text = 'Species'

# Fill rows
for ring_id, species in data:
    row_cells = table.add_row().cells
    row_cells[0].text = ring_id
    row_cells[1].text = species

doc.save('report_output.docx')

"""--------------3-------------------"""
from pypdf import PdfReader

reader = PdfReader("scientific_paper_v1.pdf")
page = reader.pages[0]

# Extract text content
raw_text = page.extract_text()

# Basic filtering
if "Parus major" in raw_text:
    print("Found mention of Kohlmeise on page 1")

"""--------------4-------------------"""
import pdfplumber

# Opening a PDF containing morphometric tables
with pdfplumber.open("data_sheet.pdf") as pdf:
  first_page = pdf.pages[0]

  # Extract table structure
  # Returns a list of lists (rows/columns)
  table_data = first_page.extract_table()

  for row in table_data:
    # Filter out None values or empty rows
    if row and row[0] != "Date":
      print(f"Date: {row[0]}, Count: {row[1]}")
