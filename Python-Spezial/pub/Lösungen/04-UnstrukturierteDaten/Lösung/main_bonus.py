import pdfplumber
from docx import Document
import re


def advanced_processing():
    # Daten-Container
    extracted_transactions = []
    applicant_name = "Max Mustermann"  # Annahme aus Teil 1, oder erneut auslesen

    print("--- Bonus: PDF Table Parsing ---")

    with pdfplumber.open("bank_statement.pdf") as pdf:
        first_page = pdf.pages[0]
        text = first_page.extract_text()

        # Wir parsen den Text Zeile für Zeile
        lines = text.split('\n')

        for line in lines:
            # Suche nach Zeilen, die mit einem Datum starten (Simple Regex)
            # Format: 2024-xx-xx
            if re.match(r'\d{4}-\d{2}-\d{2}', line):
                parts = line.split()  # Split bei Leerzeichen
                # Vorsicht: Das ist brüchig und hängt vom PDF Layout ab!
                # parts[0] = Datum, parts[1] = Text, parts[2] = Betrag
                if len(parts) >= 3:
                    tx = {
                        "date": parts[0],
                        "desc": parts[1],
                        "amount": parts[-1]
                    }
                    extracted_transactions.append(tx)

    print(f"Extrahierte Transaktionen: {len(extracted_transactions)}")
    for tx in extracted_transactions:
        print(tx)

    print("\n--- Bonus: Word Report Generierung ---")

    report = Document()
    report.add_heading('Kreditentscheidung', 0)

    # Text zusammenbauen
    p = report.add_paragraph('Basierend auf der Prüfung von ')
    p.add_run(applicant_name).bold = True
    p.add_run(' und den folgenden Transaktionen:')

    # Tabelle im Word erstellen
    table = report.add_table(rows=1, cols=3)
    # Header
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Datum'
    hdr_cells[1].text = 'Beschreibung'
    hdr_cells[2].text = 'Betrag'

    # Daten füllen
    for tx in extracted_transactions:
        row_cells = table.add_row().cells
        row_cells[0].text = tx["date"]
        row_cells[1].text = tx["desc"]
        row_cells[2].text = tx["amount"]

    report.add_paragraph('\nEntscheidung: GENEHMIGT')

    report.save('genehmigung.docx')
    print("Report 'genehmigung.docx' wurde erstellt.")


if __name__ == "__main__":
    advanced_processing()

