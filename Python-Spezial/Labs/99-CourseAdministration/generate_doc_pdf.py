from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4

# 1. Word Dokument erstellen
doc = Document()
doc.add_heading('Kreditantrag #2024-99', 0)
doc.add_paragraph('Antragsteller: Max Mustermann')
doc.add_paragraph('Kunde bittet um Kreditprüfung für Immobilienkauf.')
doc.add_paragraph('Gewünschte Summe: 500.000 EUR')
doc.save('antrag.docx')

# 2. PDF mit Tabelle erstellen
c = canvas.Canvas("bank_statement.pdf", pagesize=A4)
c.drawString(100, 800, "Fremdbank Kontoauszug - 2024")
c.drawString(100, 780, "Transaktionsübersicht:")
# Einfache visuelle Tabelle (Text-basiert für pypdf, strukturiert für pdfplumber)
y = 750
data = [["Datum", "Beschreibung", "Betrag"],
        ["2024-01-01", "Gehalt", "3500.00"],
        ["2024-01-05", "Miete", "-1200.00"],
        ["2024-01-10", "Supermarkt", "-150.50"]]

for row in data:
    line = f"{row[0]}   {row[1]}   {row[2]}"
    c.drawString(100, y, line)
    y -= 20
c.save()
print("Setup fertig: 'antrag.docx' und 'bank_statement.pdf' erstellt.")